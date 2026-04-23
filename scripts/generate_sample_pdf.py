"""Build a realistic Word document and convert it to PDF via LibreOffice.

Produces data/samples/employee_handbook.pdf — a multi-section handbook with
a proper (structured) large table that pymupdf's find_tables() can detect.

Run:
    python scripts/generate_sample_pdf.py
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SAMPLES = ROOT / "data" / "samples"
DOCX_PATH = SAMPLES / "employee_handbook.docx"
PDF_PATH = SAMPLES / "employee_handbook.pdf"

EMPLOYEE_ROWS: list[tuple[str, str, str, str, str, str]] = [
    # (name, department, role, location, start_date, salary_band)
    ("Ava Patel", "Engineering", "Staff Engineer", "Bangalore", "2019-03-11", "L6"),
    ("Mateo Rossi", "Engineering", "Senior Engineer", "Milan", "2020-06-01", "L5"),
    ("Priya Nair", "Product", "Principal PM", "Bangalore", "2018-09-23", "L7"),
    ("Liam O'Connor", "Product", "Senior PM", "Dublin", "2021-01-18", "L5"),
    ("Sakura Tanaka", "Design", "Design Lead", "Tokyo", "2017-11-02", "L6"),
    ("Daniel Kim", "Design", "Senior Designer", "Seoul", "2022-04-05", "L5"),
    ("Noor Hassan", "Data Science", "Staff DS", "Dubai", "2019-12-14", "L6"),
    ("Carlos Mendes", "Data Science", "Senior DS", "São Paulo", "2021-08-09", "L5"),
    ("Elena Volkov", "Sales", "Enterprise AE", "Berlin", "2020-02-20", "L5"),
    ("Rahul Iyer", "Sales", "Mid-Market AE", "Mumbai", "2022-05-16", "L4"),
    ("Fatima Zahra", "Sales", "SDR", "Casablanca", "2023-07-03", "L3"),
    ("Jordan Blake", "Marketing", "Head of Growth", "New York", "2018-04-30", "L7"),
    ("Mei-Lin Chen", "Marketing", "Content Lead", "Singapore", "2021-10-11", "L5"),
    ("Omar Haddad", "Customer Success", "CS Manager", "Amman", "2020-11-22", "L5"),
    ("Sofia Álvarez", "Customer Success", "CS Engineer", "Madrid", "2022-09-19", "L4"),
    ("Victor Hugo", "Finance", "Controller", "Paris", "2016-08-15", "L6"),
    ("Anika Rao", "Finance", "FP&A Analyst", "Bangalore", "2023-02-06", "L4"),
    ("Tomás Bauer", "Legal", "General Counsel", "Zurich", "2017-05-29", "L7"),
    ("Leah Okafor", "Legal", "Commercial Counsel", "Lagos", "2021-12-01", "L5"),
    ("Hiro Saito", "People Ops", "Head of People", "Tokyo", "2019-07-17", "L6"),
]

LEAVE_ROWS: list[tuple[str, str, str, str]] = [
    ("Annual", "25 days", "Accrued monthly", "Carry 5 days to next year"),
    ("Sick", "10 days", "Upfront Jan 1", "No carry-over"),
    ("Parental (primary)", "16 weeks", "After 12 months tenure", "Fully paid"),
    ("Parental (secondary)", "6 weeks", "After 12 months tenure", "Fully paid"),
    ("Bereavement", "5 days", "As needed", "Extends to immediate family"),
    ("Sabbatical", "4 weeks", "Every 5 years", "Manager approval required"),
    ("Wellness", "4 days / year", "Upfront Jan 1", "No documentation needed"),
]


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_bordered_table(doc, headers: list[str], rows: list[tuple]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    return table


def build_docx(path: Path) -> None:
    doc = Document()

    # Normal style sizing.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    _add_heading(doc, "Acme Corp Employee Handbook", 0)
    doc.add_paragraph(
        "This handbook summarizes the policies, benefits, and directory "
        "information employees most frequently ask about. It is updated "
        "quarterly and supersedes any earlier printed version."
    )

    _add_heading(doc, "1. Directory", 1)
    doc.add_paragraph(
        "The following directory lists the senior individual contributors and "
        "leads across departments. Salary bands are informational only; "
        "compensation details are covered in section 3."
    )
    _add_bordered_table(
        doc,
        headers=[
            "Name",
            "Department",
            "Role",
            "Location",
            "Start date",
            "Band",
        ],
        rows=EMPLOYEE_ROWS,
    )

    doc.add_paragraph()

    _add_heading(doc, "2. Leave & Time Off", 1)
    doc.add_paragraph(
        "Acme offers the following leave types. Unless noted otherwise, "
        "requests flow through the HRIS, with manager approval required for "
        "continuous absences longer than five working days."
    )
    _add_bordered_table(
        doc,
        headers=["Leave type", "Allotment", "Accrual rule", "Notes"],
        rows=LEAVE_ROWS,
    )

    doc.add_paragraph()

    _add_heading(doc, "3. Compensation", 1)
    doc.add_paragraph(
        "Base salary is reviewed once per year in February. Employees in band "
        "L5 and above are eligible for an annual equity refresh. The "
        "performance bonus target is 10% of base for L3–L4, 15% for L5–L6, "
        "and 20% for L7+. Payouts are made in March based on the prior "
        "calendar year."
    )

    _add_heading(doc, "4. Remote Work", 1)
    doc.add_paragraph(
        "Employees who have completed their 90-day probation may work "
        "remotely full-time or on a hybrid schedule. Core collaboration hours "
        "are 10:00 to 16:00 in the employee's local time zone. A one-time "
        "home-office stipend of $750 is available, plus $500 per year for "
        "upgrades, subject to manager approval."
    )

    _add_heading(doc, "5. Security & Compliance", 1)
    doc.add_paragraph(
        "All laptops must be enrolled in the MDM and encrypted at rest. "
        "Passwords must be at least 14 characters and rotated every 180 days. "
        "Two-factor authentication using a hardware key or an authenticator "
        "app is mandatory for production systems and source-code repositories."
    )
    doc.add_paragraph(
        "Security incidents must be reported to security@acme.example within "
        "one hour of discovery. The on-call security team acknowledges within "
        "30 minutes and coordinates the response. Post-incident review is "
        "mandatory within five business days."
    )

    _add_heading(doc, "6. Travel & Expenses", 1)
    doc.add_paragraph(
        "For quarterly on-site visits, remote employees may expense flights, "
        "lodging, and ground transport up to $1,200 per trip. Meal per-diems "
        "are $75 per day domestically and $100 per day for international "
        "travel. Receipts are required for any single item over $25."
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"Wrote {path.relative_to(ROOT)}")


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert docx to PDF via headless LibreOffice."""
    soffice = (
        shutil.which("soffice")
        or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    )
    if not Path(soffice).exists():
        print("ERROR: LibreOffice (soffice) not found on PATH", file=sys.stderr)
        sys.exit(1)

    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(pdf_path.parent),
        str(docx_path),
    ]
    subprocess.run(cmd, check=True, capture_output=True)
    produced = docx_path.with_suffix(".pdf")
    if produced != pdf_path:
        produced.replace(pdf_path)
    print(f"Wrote {pdf_path.relative_to(ROOT)}")


def main() -> None:
    build_docx(DOCX_PATH)
    convert_to_pdf(DOCX_PATH, PDF_PATH)


if __name__ == "__main__":
    main()
